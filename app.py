# flask_export_backend.py
from flask import Flask, jsonify, request, send_file, render_template
from docx import Document
import io
import random
import time
import datetime
import redis
import json

app = Flask(__name__)
r = redis.Redis(host='localhost', port=6379, decode_responses=True)

latest_data = {
    "soil_moisture": 0,
    "ph": 0,
    "latency": 0,
    "timestamp": 0
}

sensor_logs = []
base_time = datetime.datetime.now().replace(hour=0, minute=0, second=0, microsecond=0) - datetime.timedelta(days=6)

for i in range(7 * 144):
    timestamp = base_time + datetime.timedelta(minutes=10 * i)
    entry = {
        "timestamp": timestamp.isoformat(),
        "moisture": round(random.uniform(45, 70), 2),
        "ph": round(random.uniform(6.0, 7.5), 2),
        "latency": random.randint(90, 200)
    }
    sensor_logs.append(entry)
    r.rpush("sensor_logs", json.dumps(entry))

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/weekly")
def weekly():
    return render_template("weekly.html")

@app.route("/summary")
def summary():
    return render_template("summary.html")

@app.route("/settings")
def settings():
    return "<h1>Settings Page (Under Development)</h1>"

@app.route("/api/latest")
def latest():
    return jsonify(latest_data)

@app.route("/api/simulate")
def simulate():
    latest_data["soil_moisture"] = round(random.uniform(30, 70), 2)
    latest_data["ph"] = round(random.uniform(5.5, 7.5), 2)
    latest_data["latency"] = random.randint(100, 250)
    latest_data["timestamp"] = time.time()
    r.rpush("sensor_logs", json.dumps({
        "timestamp": datetime.datetime.now().isoformat(),
        "moisture": latest_data["soil_moisture"],
        "ph": latest_data["ph"],
        "latency": latest_data["latency"]
    }))
    return jsonify({"status": "data simulated", **latest_data})

@app.route("/api/control")
def control():
    state = request.args.get("state")
    if state == "on":
        print("Pompa dinyalakan (simulasi)")
    elif state == "off":
        print("Pompa dimatikan (simulasi)")
    else:
        return jsonify({"status": "unknown command"})
    return jsonify({"status": f"pompa {state}"})

@app.route("/api/weekly")
def get_weekly_summary():
    logs = [json.loads(r.lindex("sensor_logs", i)) for i in range(r.llen("sensor_logs"))]
    grouped = {}
    for entry in logs:
        date = entry["timestamp"][:10]
        if date not in grouped:
            grouped[date] = {"moisture": [], "ph": [], "latency": []}
        grouped[date]["moisture"].append(entry["moisture"])
        grouped[date]["ph"].append(entry["ph"])
        grouped[date]["latency"].append(entry["latency"])

    result = []
    for i, (date, values) in enumerate(sorted(grouped.items())):
        result.append({
            "hari": f"Hari {i + 1}",
            "moisture": round(sum(values["moisture"]) / len(values["moisture"]), 2),
            "ph": round(sum(values["ph"]) / len(values["ph"]), 2),
            "latency": round(sum(values["latency"]) / len(values["latency"]), 2)
        })

    return jsonify(result)

@app.route("/export/docx")
def export_docx():
    logs = [json.loads(r.lindex("sensor_logs", i)) for i in range(r.llen("sensor_logs"))]
    doc = Document()
    doc.add_heading("Laporan Mingguan Smart Irrigation", 0)
    doc.add_paragraph(f"Tanggal: {datetime.date.today().isoformat()}")
    doc.add_paragraph("\nRangkuman Data Mingguan:")
    doc.add_paragraph(f"Soil Moisture (avg): {sum([d['moisture'] for d in logs]) / len(logs):.2f}%")
    doc.add_paragraph(f"pH Tanah (avg): {sum([d['ph'] for d in logs]) / len(logs):.2f}")
    doc.add_paragraph(f"Latency (avg): {sum([d['latency'] for d in logs]) / len(logs):.2f} ms")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="laporan_mingguan.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.route("/export/weekly.csv")
def export_weekly_csv():
    logs = [json.loads(r.lindex("sensor_logs", i)) for i in range(r.llen("sensor_logs"))]
    output = io.StringIO()
    output.write("Hari,Soil Moisture (%),pH Tanah,Latency (ms)\n")
    grouped = {}
    for entry in logs:
        date = entry["timestamp"][:10]
        if date not in grouped:
            grouped[date] = {"moisture": [], "ph": [], "latency": []}
        grouped[date]["moisture"].append(entry["moisture"])
        grouped[date]["ph"].append(entry["ph"])
        grouped[date]["latency"].append(entry["latency"])

    for i, (date, values) in enumerate(sorted(grouped.items())):
        moisture_avg = sum(values["moisture"]) / len(values["moisture"])
        ph_avg = sum(values["ph"]) / len(values["ph"])
        latency_avg = sum(values["latency"]) / len(values["latency"])
        output.write(f"Hari {i + 1},{moisture_avg:.2f},{ph_avg:.2f},{latency_avg:.2f}\n")

    output.seek(0)
    return send_file(io.BytesIO(output.read().encode()), as_attachment=True, download_name="laporan_mingguan.csv", mimetype="text/csv")

@app.route("/export/daily.docx")
def export_daily_docx():
    doc = Document()
    doc.add_heading("Laporan Harian Smart Irrigation", 0)
    doc.add_paragraph(f"Tanggal: {datetime.date.today().isoformat()}")
    doc.add_paragraph(f"Soil Moisture: {latest_data['soil_moisture']}%")
    doc.add_paragraph(f"pH Tanah: {latest_data['ph']}")
    doc.add_paragraph(f"Latency: {latest_data['latency']} ms")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="laporan_harian.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.route("/export/daily.csv")
def export_daily_csv():
    output = io.StringIO()
    output.write("Tanggal,Soil Moisture (%),pH Tanah,Latency (ms)\n")
    output.write(f"{datetime.date.today().isoformat()},{latest_data['soil_moisture']},{latest_data['ph']},{latest_data['latency']}\n")
    output.seek(0)
    return send_file(io.BytesIO(output.read().encode()), as_attachment=True, download_name="laporan_harian.csv", mimetype="text/csv")

if __name__ == '__main__':
    app.run(debug=True, port=5000)
